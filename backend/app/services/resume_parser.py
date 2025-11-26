"""
Resume parsing service for extracting text from PDF, DOCX, and TXT files.
Supports multiple parsing strategies with fallbacks for robustness.
"""
import io
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document

from app.core.logging_config import get_logger

logger = get_logger(__name__)


class ResumeParseError(Exception):
    """Exception raised when resume parsing fails."""
    pass


class ResumeParser:
    """Service for parsing resumes in various formats."""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".txt"}

    @staticmethod
    def parse(file_content: bytes, filename: str) -> str:
        """
        Parse resume from file content.

        Args:
            file_content: Raw file bytes
            filename: Original filename with extension

        Returns:
            Extracted text content

        Raises:
            ResumeParseError: If parsing fails
        """
        file_extension = Path(filename).suffix.lower()

        if file_extension not in ResumeParser.SUPPORTED_FORMATS:
            raise ResumeParseError(
                f"Unsupported file format: {file_extension}. "
                f"Supported formats: {', '.join(ResumeParser.SUPPORTED_FORMATS)}"
            )

        logger.info("Parsing resume", filename=filename, format=file_extension)

        try:
            if file_extension == ".pdf":
                text = ResumeParser._parse_pdf(file_content)
            elif file_extension == ".docx":
                text = ResumeParser._parse_docx(file_content)
            elif file_extension == ".txt":
                text = ResumeParser._parse_txt(file_content)
            else:
                raise ResumeParseError(f"Unsupported format: {file_extension}")

            if not text or not text.strip():
                raise ResumeParseError("No text content extracted from file")

            logger.info(
                "Resume parsed successfully",
                filename=filename,
                text_length=len(text)
            )

            return text.strip()

        except Exception as e:
            logger.error("Resume parsing failed", filename=filename, error=str(e))
            raise ResumeParseError(f"Failed to parse {filename}: {str(e)}")

    @staticmethod
    def _parse_pdf(file_content: bytes) -> str:
        """
        Parse PDF file using PyMuPDF.

        Args:
            file_content: PDF file bytes

        Returns:
            Extracted text
        """
        try:
            # Open PDF from bytes
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            text_parts = []

            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text_parts.append(page.get_text())

            pdf_document.close()

            text = "\n\n".join(text_parts)
            return text

        except Exception as e:
            logger.error("PDF parsing error", error=str(e))
            raise ResumeParseError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def _parse_docx(file_content: bytes) -> str:
        """
        Parse DOCX file using python-docx.

        Args:
            file_content: DOCX file bytes

        Returns:
            Extracted text
        """
        try:
            # Open DOCX from bytes
            doc = Document(io.BytesIO(file_content))
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            text = "\n".join(text_parts)
            return text

        except Exception as e:
            logger.error("DOCX parsing error", error=str(e))
            raise ResumeParseError(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    def _parse_txt(file_content: bytes) -> str:
        """
        Parse TXT file with encoding detection.

        Args:
            file_content: TXT file bytes

        Returns:
            Extracted text
        """
        try:
            # Try UTF-8 first
            try:
                text = file_content.decode("utf-8")
            except UnicodeDecodeError:
                # Fallback to latin-1
                try:
                    text = file_content.decode("latin-1")
                except UnicodeDecodeError:
                    # Last resort: ignore errors
                    text = file_content.decode("utf-8", errors="ignore")

            return text

        except Exception as e:
            logger.error("TXT parsing error", error=str(e))
            raise ResumeParseError(f"Failed to parse TXT: {str(e)}")


class ResumeAnalyzer:
    """Service for analyzing resume content using LLMs."""

    ANALYSIS_PROMPT = """
You are an expert resume analyzer. Analyze the following resume and extract key information in a structured format.

Resume:
{resume_text}

Extract and provide:
1. **Skills**: List all technical skills, soft skills, and competencies
2. **Experience**: Summarize work experience with years and key roles
3. **Education**: List educational qualifications
4. **Keywords**: Extract important keywords that would be relevant for job matching
5. **Summary**: A brief professional summary (2-3 sentences)

Format your response as JSON with the following structure:
{{
    "skills": ["skill1", "skill2", ...],
    "experience": {{
        "total_years": <number>,
        "roles": ["role1", "role2", ...]
    }},
    "education": ["degree1", "degree2", ...],
    "keywords": ["keyword1", "keyword2", ...],
    "summary": "Professional summary text"
}}
"""

    def __init__(self, llm_client):
        """
        Initialize analyzer with an LLM client.

        Args:
            llm_client: Instance of BaseLLMClient
        """
        self.llm_client = llm_client

    async def analyze(self, resume_text: str) -> dict:
        """
        Analyze resume text using LLM.

        Args:
            resume_text: Extracted resume text

        Returns:
            Structured analysis results
        """
        try:
            logger.info("Analyzing resume with LLM")

            prompt = self.ANALYSIS_PROMPT.format(resume_text=resume_text)

            response = await self.llm_client.generate(
                prompt=prompt,
                temperature=0.3,
                max_tokens=2048
            )

            logger.info("Resume analysis completed")

            # Try to parse JSON response
            import json
            try:
                # Extract JSON from response (LLM might wrap it in markdown)
                content = response.content
                if "```json" in content:
                    content = content.split("```json")[1].split("```")[0]
                elif "```" in content:
                    content = content.split("```")[1].split("```")[0]

                analysis_data = json.loads(content.strip())

            except json.JSONDecodeError:
                logger.warning("Failed to parse JSON, returning raw content")
                analysis_data = {"raw_analysis": response.content}

            # Add metadata
            analysis_data["_metadata"] = {
                "model": response.model,
                "provider": response.provider,
                "tokens_used": response.tokens_used,
                "cost_estimate": response.cost_estimate
            }

            return analysis_data

        except Exception as e:
            logger.error("Resume analysis failed", error=str(e))
            raise
