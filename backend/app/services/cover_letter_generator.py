"""
Cover letter generator for creating tailored cover letters.
Uses LLM to generate professional cover letters based on resume and job description.
"""
from typing import Dict, List, Any, Optional
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from datetime import datetime

from app.core.logging_config import get_logger
from app.core.llm_providers import BaseLLMClient

logger = get_logger(__name__)


class CoverLetterGenerator:
    """Generate tailored cover letters."""

    def __init__(self, llm_client: BaseLLMClient):
        self.llm_client = llm_client

    async def generate(
        self,
        resume_text: str,
        job_description: str,
        job_title: str,
        company: str,
        candidate_name: Optional[str] = None,
        tone: str = "professional"
    ) -> Dict[str, Any]:
        """
        Generate a tailored cover letter.

        Args:
            resume_text: Candidate's resume text
            job_description: Job description
            job_title: Job title
            company: Company name
            candidate_name: Optional candidate name
            tone: Tone style (professional, enthusiastic, formal)

        Returns:
            Dictionary with cover letter text and metadata
        """
        try:
            prompt = self._build_prompt(
                resume_text=resume_text,
                job_description=job_description,
                job_title=job_title,
                company=company,
                candidate_name=candidate_name,
                tone=tone
            )

            response = await self.llm_client.generate(prompt)

            # Extract name from resume if not provided
            if not candidate_name:
                candidate_name = self._extract_name(resume_text)

            result = {
                "cover_letter": response,
                "candidate_name": candidate_name,
                "job_title": job_title,
                "company": company,
                "tone": tone
            }

            logger.info("Cover letter generated successfully")
            return result

        except Exception as e:
            logger.error("Failed to generate cover letter", error=str(e))
            raise

    def _build_prompt(
        self,
        resume_text: str,
        job_description: str,
        job_title: str,
        company: str,
        candidate_name: Optional[str],
        tone: str
    ) -> str:
        """Build the LLM prompt for cover letter generation."""

        tone_guidance = {
            "professional": "Use a professional, confident tone. Be direct and results-oriented.",
            "enthusiastic": "Use an enthusiastic, passionate tone while maintaining professionalism. Show genuine excitement.",
            "formal": "Use a formal, traditional business tone. Be respectful and diplomatic."
        }

        tone_instruction = tone_guidance.get(tone, tone_guidance["professional"])

        name_instruction = f"The candidate's name is {candidate_name}." if candidate_name else "Extract the candidate's name from the resume."

        prompt = f"""You are an expert career advisor and professional writer. Write a compelling, tailored cover letter.

**Job Information:**
Position: {job_title}
Company: {company}

**Job Description:**
{job_description}

**Candidate's Resume:**
{resume_text}

**Instructions:**
1. {name_instruction}
2. {tone_instruction}
3. Structure the cover letter with:
   - Opening paragraph: Express interest and briefly state why you're a strong fit
   - 2-3 body paragraphs: Highlight specific, relevant accomplishments from the resume that align with the job requirements
   - Closing paragraph: Express enthusiasm, mention next steps, and thank them

4. Be specific: Reference actual projects, achievements, and experiences from the resume
5. Show you researched the company: Reference the specific role and requirements
6. Keep it concise: Aim for 3-4 paragraphs, about 250-350 words
7. Make it unique: Avoid generic phrases like "I am writing to apply for..."
8. Use quantifiable achievements when possible
9. Show personality while maintaining professionalism

**Format:**
Write ONLY the body of the cover letter (no date, address, or signature block - just the letter content).
Start with the opening paragraph immediately.

Do NOT include:
- Today's date
- Addresses
- "Sincerely," or signature
- Placeholder text like [Your Name]

These will be added separately in the formatted document.

Write the cover letter now:
"""

        return prompt

    def _extract_name(self, resume_text: str) -> str:
        """Extract candidate name from resume (simple heuristic)."""
        lines = resume_text.strip().split('\n')

        # Usually the name is in the first few lines
        for line in lines[:5]:
            line = line.strip()
            # Look for a line with 2-4 capitalized words (likely a name)
            words = line.split()
            if 2 <= len(words) <= 4:
                if all(word[0].isupper() for word in words if word):
                    return line

        return "Candidate"

    @staticmethod
    def create_docx(
        cover_letter_text: str,
        candidate_name: str,
        candidate_email: Optional[str] = None,
        candidate_phone: Optional[str] = None,
        company: str = "Hiring Manager",
        job_title: str = "Position"
    ) -> BytesIO:
        """
        Create a professionally formatted cover letter DOCX.

        Args:
            cover_letter_text: Cover letter body text
            candidate_name: Candidate's name
            candidate_email: Optional email
            candidate_phone: Optional phone
            company: Company name
            job_title: Job title

        Returns:
            BytesIO object containing the DOCX file
        """
        try:
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Header: Candidate contact info
            header = doc.add_paragraph()
            header_run = header.add_run(candidate_name)
            header_run.bold = True
            header_run.font.size = Pt(14)

            if candidate_email or candidate_phone:
                contact_para = doc.add_paragraph()
                contact_parts = []
                if candidate_email:
                    contact_parts.append(candidate_email)
                if candidate_phone:
                    contact_parts.append(candidate_phone)
                contact_para.add_run(" | ".join(contact_parts)).font.size = Pt(10)

            doc.add_paragraph()  # Spacing

            # Date
            date_para = doc.add_paragraph()
            date_para.add_run(datetime.now().strftime("%B %d, %Y")).font.size = Pt(11)
            doc.add_paragraph()  # Spacing

            # Recipient
            recipient = doc.add_paragraph()
            recipient.add_run(f"Hiring Manager\n{company}").font.size = Pt(11)
            doc.add_paragraph()  # Spacing

            # Salutation
            salutation = doc.add_paragraph()
            salutation.add_run("Dear Hiring Manager,").font.size = Pt(11)
            doc.add_paragraph()  # Spacing

            # Body paragraphs
            paragraphs = cover_letter_text.strip().split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(para_text.strip())
                    run.font.size = Pt(11)
                    p.paragraph_format.space_after = Pt(12)
                    p.paragraph_format.line_spacing = 1.15

            # Closing
            doc.add_paragraph()
            closing = doc.add_paragraph()
            closing.add_run("Sincerely,").font.size = Pt(11)

            doc.add_paragraph()
            signature = doc.add_paragraph()
            signature.add_run(candidate_name).font.size = Pt(11)

            # Save to BytesIO
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            logger.info("Cover letter DOCX created successfully")
            return file_stream

        except Exception as e:
            logger.error("Failed to create cover letter DOCX", error=str(e))
            raise

    @staticmethod
    def create_pdf(
        cover_letter_text: str,
        candidate_name: str,
        candidate_email: Optional[str] = None,
        candidate_phone: Optional[str] = None,
        company: str = "Hiring Manager",
        job_title: str = "Position"
    ) -> BytesIO:
        """
        Create a professionally formatted cover letter PDF.

        Args:
            cover_letter_text: Cover letter body text
            candidate_name: Candidate's name
            candidate_email: Optional email
            candidate_phone: Optional phone
            company: Company name
            job_title: Job title

        Returns:
            BytesIO object containing the PDF file
        """
        try:
            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=letter,
                rightMargin=1*inch,
                leftMargin=1*inch,
                topMargin=1*inch,
                bottomMargin=1*inch
            )

            styles = getSampleStyleSheet()

            # Custom styles
            name_style = ParagraphStyle(
                'Name',
                parent=styles['Normal'],
                fontSize=14,
                fontName='Helvetica-Bold',
                spaceAfter=6
            )

            contact_style = ParagraphStyle(
                'Contact',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=12
            )

            body_style = ParagraphStyle(
                'Body',
                parent=styles['Normal'],
                fontSize=11,
                leading=14,
                spaceAfter=12,
                alignment=TA_JUSTIFY
            )

            story = []

            # Header
            story.append(Paragraph(candidate_name, name_style))

            if candidate_email or candidate_phone:
                contact_parts = []
                if candidate_email:
                    contact_parts.append(candidate_email)
                if candidate_phone:
                    contact_parts.append(candidate_phone)
                story.append(Paragraph(" | ".join(contact_parts), contact_style))

            story.append(Spacer(1, 0.2*inch))

            # Date
            story.append(Paragraph(datetime.now().strftime("%B %d, %Y"), styles['Normal']))
            story.append(Spacer(1, 0.2*inch))

            # Recipient
            story.append(Paragraph(f"Hiring Manager<br/>{company}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))

            # Salutation
            story.append(Paragraph("Dear Hiring Manager,", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))

            # Body
            paragraphs = cover_letter_text.strip().split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    story.append(Paragraph(para_text.strip(), body_style))

            # Closing
            story.append(Spacer(1, 0.2*inch))
            story.append(Paragraph("Sincerely,", styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
            story.append(Paragraph(candidate_name, styles['Normal']))

            doc.build(story)
            pdf_buffer.seek(0)

            logger.info("Cover letter PDF created successfully")
            return pdf_buffer

        except Exception as e:
            logger.error("Failed to create cover letter PDF", error=str(e))
            raise
