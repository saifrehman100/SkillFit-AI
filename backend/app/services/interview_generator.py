"""
Interview question generator for preparing candidates based on resume and job match.
Uses LLM to generate tailored interview questions and talking points.
"""
from typing import Dict, List, Any, Optional
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

from app.core.logging_config import get_logger
from app.core.llm_providers import BaseLLMClient

logger = get_logger(__name__)


class InterviewGenerator:
    """Generate tailored interview preparation materials."""

    def __init__(self, llm_client: BaseLLMClient):
        self.llm_client = llm_client

    async def generate_questions(
        self,
        resume_text: str,
        job_description: str,
        job_title: str = "the position",
        company: str = "the company",
        match_score: Optional[float] = None,
        missing_skills: Optional[List[str]] = None,
        recommendations: Optional[List[Any]] = None
    ) -> Dict[str, Any]:
        """
        Generate interview questions and preparation tips.

        Args:
            resume_text: Candidate's resume text
            job_description: Job description
            job_title: Job title
            company: Company name
            match_score: Optional match score
            missing_skills: Optional list of missing skills
            recommendations: Optional improvement recommendations

        Returns:
            Dictionary with technical_questions, behavioral_questions,
            gap_questions, and talking_points
        """
        try:
            prompt = self._build_prompt(
                resume_text=resume_text,
                job_description=job_description,
                job_title=job_title,
                company=company,
                match_score=match_score,
                missing_skills=missing_skills,
                recommendations=recommendations
            )

            response = await self.llm_client.generate(prompt)
            result = self._parse_response(response.content)

            logger.info("Interview questions generated successfully")
            return result

        except Exception as e:
            logger.error("Failed to generate interview questions", error=str(e))
            raise

    def _build_prompt(
        self,
        resume_text: str,
        job_description: str,
        job_title: str,
        company: str,
        match_score: Optional[float],
        missing_skills: Optional[List[str]],
        recommendations: Optional[List[Any]]
    ) -> str:
        """Build the LLM prompt for interview question generation."""

        prompt = f"""You are an expert interview coach preparing a candidate for a job interview.

**Job Information:**
Position: {job_title}
Company: {company}

**Job Description:**
{job_description}

**Candidate's Resume:**
{resume_text}
"""

        if match_score is not None:
            prompt += f"\n**Match Score:** {match_score}/100"

        if missing_skills:
            prompt += f"\n**Skills to Address:**\n" + "\n".join(f"- {skill}" for skill in missing_skills[:5])

        prompt += """

Generate a comprehensive interview preparation guide with the following sections:

1. **TECHNICAL QUESTIONS** (5-7 questions)
   - Questions the interviewer is likely to ask based on job requirements
   - Include both questions and suggested answer approaches
   - Focus on skills mentioned in the job description

2. **BEHAVIORAL QUESTIONS** (4-6 questions)
   - STAR method questions related to experiences on the resume
   - Include suggested examples from the candidate's background
   - Cover leadership, teamwork, problem-solving, conflict resolution

3. **GAP/WEAKNESS QUESTIONS** (2-4 questions)
   - Address any gaps, missing skills, or potential concerns
   - Provide honest, positive ways to address these
   - Turn weaknesses into growth opportunities

4. **KEY TALKING POINTS** (5-7 points)
   - Most impressive achievements to highlight
   - How candidate's experience aligns with role
   - Unique value propositions
   - Questions candidate should ask the interviewer

Format your response EXACTLY as follows:

## TECHNICAL QUESTIONS

Q1: [Question text]
A: [Suggested approach to answer, referring to specific resume experiences]

Q2: [Question text]
A: [Suggested approach]

[Continue for all technical questions]

## BEHAVIORAL QUESTIONS

Q1: [Question text]
A: [Suggested STAR example from resume]

Q2: [Question text]
A: [Suggested example]

[Continue for all behavioral questions]

## GAP/WEAKNESS QUESTIONS

Q1: [Question text]
A: [Honest, positive way to address this]

Q2: [Question text]
A: [How to reframe as growth opportunity]

[Continue for all gap questions]

## KEY TALKING POINTS

1. [Talking point with specific details]
2. [Talking point]
3. [Talking point]
[Continue for all talking points]

Be specific, actionable, and tailored to this exact job and candidate combination.
"""

        return prompt

    def _parse_response(self, response: str) -> Dict[str, Any]:
        """Parse LLM response into structured format."""

        sections = {
            "technical_questions": [],
            "behavioral_questions": [],
            "gap_questions": [],
            "talking_points": [],
            "raw_text": response
        }

        current_section = None
        current_question = None
        current_answer = ""

        lines = response.split('\n')

        for line in lines:
            line_stripped = line.strip()

            # Detect section headers
            if "TECHNICAL QUESTIONS" in line_stripped.upper():
                current_section = "technical_questions"
                continue
            elif "BEHAVIORAL QUESTIONS" in line_stripped.upper():
                current_section = "behavioral_questions"
                continue
            elif "GAP" in line_stripped.upper() and "QUESTIONS" in line_stripped.upper():
                current_section = "gap_questions"
                continue
            elif "TALKING POINTS" in line_stripped.upper() or "KEY POINTS" in line_stripped.upper():
                current_section = "talking_points"
                continue

            if not current_section or not line_stripped:
                continue

            # Parse talking points (numbered lists)
            if current_section == "talking_points":
                # Match numbered or bulleted items
                if line_stripped[0].isdigit() or line_stripped.startswith('-') or line_stripped.startswith('•'):
                    # Remove number/bullet
                    point = line_stripped.lstrip('0123456789.-•').strip()
                    if point:
                        sections["talking_points"].append(point)

            # Parse Q&A format
            else:
                if line_stripped.startswith('Q') and ':' in line_stripped:
                    # Save previous Q&A if exists
                    if current_question and current_answer:
                        sections[current_section].append({
                            "question": current_question,
                            "answer": current_answer.strip()
                        })

                    # Start new question
                    current_question = line_stripped.split(':', 1)[1].strip()
                    current_answer = ""

                elif line_stripped.startswith('A') and ':' in line_stripped and current_question:
                    # Start answer
                    current_answer = line_stripped.split(':', 1)[1].strip()

                elif current_question and current_answer:
                    # Continue answer
                    current_answer += " " + line_stripped

        # Save last Q&A
        if current_section in ["technical_questions", "behavioral_questions", "gap_questions"]:
            if current_question and current_answer:
                sections[current_section].append({
                    "question": current_question,
                    "answer": current_answer.strip()
                })

        return sections

    @staticmethod
    def create_docx(
        interview_data: Dict[str, Any],
        job_title: str = "Position",
        company: str = "Company"
    ) -> BytesIO:
        """
        Create a professionally formatted DOCX interview prep document.

        Args:
            interview_data: Generated interview questions and talking points
            job_title: Job title
            company: Company name

        Returns:
            BytesIO object containing the DOCX file
        """
        try:
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Title
            title = doc.add_paragraph()
            run = title.add_run(f"Interview Preparation Guide")
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(31, 78, 120)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Job info
            job_info = doc.add_paragraph()
            job_info.add_run(f"{job_title} at {company}").font.size = Pt(12)
            job_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            # Technical Questions
            if interview_data.get("technical_questions"):
                heading = doc.add_paragraph()
                run = heading.add_run("Technical Questions")
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(31, 78, 120)
                doc.add_paragraph()

                for i, qa in enumerate(interview_data["technical_questions"], 1):
                    # Question
                    q_para = doc.add_paragraph()
                    q_run = q_para.add_run(f"Q{i}: {qa['question']}")
                    q_run.bold = True
                    q_run.font.size = Pt(11)

                    # Answer
                    a_para = doc.add_paragraph()
                    a_para.add_run(f"A: {qa['answer']}").font.size = Pt(10)
                    a_para.paragraph_format.left_indent = Inches(0.25)
                    doc.add_paragraph()

            # Behavioral Questions
            if interview_data.get("behavioral_questions"):
                heading = doc.add_paragraph()
                run = heading.add_run("Behavioral Questions")
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(31, 78, 120)
                doc.add_paragraph()

                for i, qa in enumerate(interview_data["behavioral_questions"], 1):
                    q_para = doc.add_paragraph()
                    q_run = q_para.add_run(f"Q{i}: {qa['question']}")
                    q_run.bold = True
                    q_run.font.size = Pt(11)

                    a_para = doc.add_paragraph()
                    a_para.add_run(f"A: {qa['answer']}").font.size = Pt(10)
                    a_para.paragraph_format.left_indent = Inches(0.25)
                    doc.add_paragraph()

            # Gap/Weakness Questions
            if interview_data.get("gap_questions"):
                heading = doc.add_paragraph()
                run = heading.add_run("Addressing Gaps & Weaknesses")
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(31, 78, 120)
                doc.add_paragraph()

                for i, qa in enumerate(interview_data["gap_questions"], 1):
                    q_para = doc.add_paragraph()
                    q_run = q_para.add_run(f"Q{i}: {qa['question']}")
                    q_run.bold = True
                    q_run.font.size = Pt(11)

                    a_para = doc.add_paragraph()
                    a_para.add_run(f"A: {qa['answer']}").font.size = Pt(10)
                    a_para.paragraph_format.left_indent = Inches(0.25)
                    doc.add_paragraph()

            # Key Talking Points
            if interview_data.get("talking_points"):
                heading = doc.add_paragraph()
                run = heading.add_run("Key Talking Points")
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(31, 78, 120)
                doc.add_paragraph()

                for point in interview_data["talking_points"]:
                    p = doc.add_paragraph(point, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.25)

            # Save to BytesIO
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            logger.info("Interview prep DOCX created successfully")
            return file_stream

        except Exception as e:
            logger.error("Failed to create interview DOCX", error=str(e))
            raise

    @staticmethod
    def create_pdf(
        interview_data: Dict[str, Any],
        job_title: str = "Position",
        company: str = "Company"
    ) -> BytesIO:
        """
        Create a professionally formatted PDF interview prep document.

        Args:
            interview_data: Generated interview questions and talking points
            job_title: Job title
            company: Company name

        Returns:
            BytesIO object containing the PDF file
        """
        try:
            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=letter,
                rightMargin=1*inch,
                leftMargin=1*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )

            styles = getSampleStyleSheet()

            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=RGBColor(31/255, 78/255, 120/255),
                spaceAfter=12,
                alignment=1  # Center
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=RGBColor(31/255, 78/255, 120/255),
                spaceAfter=12,
                spaceBefore=12
            )

            question_style = ParagraphStyle(
                'Question',
                parent=styles['Normal'],
                fontSize=11,
                fontName='Helvetica-Bold',
                spaceAfter=6
            )

            answer_style = ParagraphStyle(
                'Answer',
                parent=styles['Normal'],
                fontSize=10,
                leftIndent=20,
                spaceAfter=12
            )

            story = []

            # Title
            story.append(Paragraph(f"Interview Preparation Guide", title_style))
            story.append(Paragraph(f"{job_title} at {company}", styles['Normal']))
            story.append(Spacer(1, 0.3*inch))

            # Technical Questions
            if interview_data.get("technical_questions"):
                story.append(Paragraph("Technical Questions", heading_style))

                for i, qa in enumerate(interview_data["technical_questions"], 1):
                    story.append(Paragraph(f"Q{i}: {qa['question']}", question_style))
                    story.append(Paragraph(f"A: {qa['answer']}", answer_style))

            # Behavioral Questions
            if interview_data.get("behavioral_questions"):
                story.append(Paragraph("Behavioral Questions", heading_style))

                for i, qa in enumerate(interview_data["behavioral_questions"], 1):
                    story.append(Paragraph(f"Q{i}: {qa['question']}", question_style))
                    story.append(Paragraph(f"A: {qa['answer']}", answer_style))

            # Gap Questions
            if interview_data.get("gap_questions"):
                story.append(Paragraph("Addressing Gaps & Weaknesses", heading_style))

                for i, qa in enumerate(interview_data["gap_questions"], 1):
                    story.append(Paragraph(f"Q{i}: {qa['question']}", question_style))
                    story.append(Paragraph(f"A: {qa['answer']}", answer_style))

            # Talking Points
            if interview_data.get("talking_points"):
                story.append(Paragraph("Key Talking Points", heading_style))

                for point in interview_data["talking_points"]:
                    story.append(Paragraph(f"• {point}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))

            doc.build(story)
            pdf_buffer.seek(0)

            logger.info("Interview prep PDF created successfully")
            return pdf_buffer

        except Exception as e:
            logger.error("Failed to create interview PDF", error=str(e))
            raise
