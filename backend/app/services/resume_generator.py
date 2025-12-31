"""
Resume document generator for creating downloadable resume files.
Converts resume text to formatted DOCX files.
"""
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.logging_config import get_logger

logger = get_logger(__name__)


class ResumeGenerator:
    """Generate formatted resume documents."""

    @staticmethod
    def create_docx(resume_text: str, filename: str = "resume.docx") -> BytesIO:
        """
        Convert resume text to a formatted DOCX file.

        Args:
            resume_text: Plain text resume content
            filename: Optional filename for the document

        Returns:
            BytesIO object containing the DOCX file
        """
        try:
            doc = Document()

            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Parse the resume text and add formatted content
            lines = resume_text.split('\n')

            for line in lines:
                line = line.strip()

                if not line:
                    # Empty line - add spacing
                    doc.add_paragraph()
                    continue

                # Detect section headers (usually all caps or ending with colon)
                is_header = (
                    line.isupper() or
                    line.endswith(':') or
                    any(keyword in line.upper() for keyword in [
                        'SUMMARY', 'EXPERIENCE', 'EDUCATION', 'SKILLS',
                        'PROJECTS', 'CERTIFICATIONS', 'CONTACT'
                    ])
                )

                if is_header:
                    # Add section header with bold formatting
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # Regular text
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(6)

                    # Check if it's a bullet point
                    if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                        p.paragraph_format.left_indent = Inches(0.25)

            # Save to BytesIO
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            logger.info("Resume DOCX created successfully")
            return file_stream

        except Exception as e:
            logger.error("Failed to create DOCX", error=str(e))
            raise

    @staticmethod
    def create_professional_docx(
        resume_text: str,
        candidate_name: str = None,
        filename: str = "resume.docx"
    ) -> BytesIO:
        """
        Create a professionally formatted resume DOCX.

        Args:
            resume_text: Plain text resume content
            candidate_name: Optional candidate name for header
            filename: Optional filename

        Returns:
            BytesIO object containing the DOCX file
        """
        try:
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Add candidate name as header if provided
            if candidate_name:
                header = doc.add_paragraph()
                run = header.add_run(candidate_name)
                run.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(31, 78, 120)  # Professional blue
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Spacing

            # Process resume text
            current_section = None
            lines = resume_text.split('\n')

            for line in lines:
                line = line.strip()

                if not line:
                    continue

                # Detect section headers
                is_major_header = any(keyword in line.upper() for keyword in [
                    'PROFESSIONAL SUMMARY', 'SUMMARY', 'OBJECTIVE',
                    'WORK EXPERIENCE', 'EXPERIENCE', 'PROFESSIONAL EXPERIENCE',
                    'EDUCATION', 'SKILLS', 'TECHNICAL SKILLS',
                    'PROJECTS', 'CERTIFICATIONS', 'ACHIEVEMENTS'
                ])

                if is_major_header:
                    # Major section header
                    doc.add_paragraph()  # Add space before section
                    p = doc.add_paragraph()
                    run = p.add_run(line.upper())
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(31, 78, 120)

                    # Add bottom border to section header
                    p.paragraph_format.space_after = Pt(6)
                    current_section = line

                elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    # Bullet point
                    p = doc.add_paragraph(line, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(3)

                else:
                    # Regular paragraph or job title/company
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(6)

                    # Make job titles bold if in experience section
                    if current_section and 'EXPERIENCE' in current_section.upper():
                        if '|' in line or '–' in line or '-' in line:
                            # Likely a job title with company/date
                            run = p.runs[0]
                            run.bold = True

            # Save to BytesIO
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            logger.info("Professional resume DOCX created successfully")
            return file_stream

        except Exception as e:
            logger.error("Failed to create professional DOCX", error=str(e))
            raise
